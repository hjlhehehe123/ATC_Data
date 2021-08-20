from django.shortcuts import render
from docx import Document
import time

from xunliankaohe import models
from jianchabaogao.models import jianchabaogao
from atcInfo.models import info_02


# Create your views here.



def getsheet(request):

    selLocation1 = request.POST.get('selLocation1')
    selLocation2 = request.POST.get('selLocation2')
    selLocation3 = request.POST.get('selLocation3')

    if selLocation1 == "0" and selLocation2 == "6" and selLocation3 == "8":
        return render(request, 'jichangfxkhb.html')
    if selLocation1 == "1" and selLocation2 == "6" and selLocation3 == "8":
        return render(request, 'APPfxkhb.html')
    if selLocation1 == "2" and selLocation2 == "6" and selLocation3 == "8":
        return render(request, 'ACCfxkhb.html')
    if selLocation1 == "0" and selLocation2 == "4" and selLocation3 == "8":
        return render(request, 'jichanggqkhb.html')
    if selLocation1 == "1" and selLocation2 == "4" and selLocation3 == "8":
        return render(request, 'APPgqkhb.html')
    if selLocation1 == "2" and selLocation2 == "4" and selLocation3 == "8":
        return render(request, 'ACCgqkhb.html')
    if selLocation1 == "0" and selLocation2 == "5" and selLocation3 == "8":
        return render(request, 'jichangzgkhb.html')
    if selLocation1 == "1" and selLocation2 == "5" and selLocation3 == "8":
        return render(request, 'APPzgkhb.html')
    if selLocation1 == "2" and selLocation2 == "5" and selLocation3 == "8":
        return render(request, 'ACCzgkhb.html')


def addjichangfxkhb(request):  # 获取增加机场复训考核表记录的数据并保存

    textlist = []
    listnum=[11,1,1,1,1,1,1,2,1,2,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1]

    datan = request.POST.get('data01')
    datas = request.POST.get('data02')

    for i in range(1,8):
        name = 'data0' + str(i)
        a=request.POST.get(name)
        textlist.append(a)

    data08 = request.POST.get('data08')
    if data08 == '0':
        data08 = '初级'
    if data08 == '1':
        data08 = '中级'
    if data08 == '2':
        data08 = '高级'
    textlist.append(data08)
    data09 = request.POST.get('data09')
    textlist.append(data09)

    for i in range(10,40):
        name = 'data' + str(i)
        a=request.POST.get(name)
        b=int(a)
        c = b * listnum[i-10]
        b=str(b)
        c=str(c)
        textlist.append(b)
        textlist.append(c)

    data40 = request.POST.get('jgcs')
    data41 = request.POST.get('jgkf')
    data42 = request.POST.get('datajgz')
    data43 = request.POST.get('xtcs')
    data44 = request.POST.get('xtkf')
    data45 = request.POST.get('dataxtz')
    data46 = request.POST.get('yscs')
    data47 = request.POST.get('yskf')
    data48 = request.POST.get('dataysz')
    data49 = request.POST.get('lkcs')
    data50 = request.POST.get('lkkf')
    data51 = request.POST.get('datalkz')
    data52 = request.POST.get('tqcs')
    data53 = request.POST.get('tqkf')
    data54 = request.POST.get('datatqz')
    data55 = request.POST.get('ldcs')
    data56 = request.POST.get('ldkf')
    data57 = request.POST.get('dataldz')
    data58 = request.POST.get('qtkf')
    data59 = request.POST.get('data41')
    if data59=='':
       data59="无"

    textlist.append(data59)
    dataqtkf = request.POST.get('data40')
    textlist.append(dataqtkf)

    data61 = request.POST.get('qkjl')
    textlist.append(data61)
    data62 = request.POST.get('zwpj')
    textlist.append(data62)
    data63 = request.POST.get('jypy')
    textlist.append(data63)
    data64 = request.POST.get('jg')
    if data64 == '0':
        data64 = '通过'
    if data64 == '1':
        data64 = '不通过'
    if data64 == '2':
        data64 = '待定'
    data65 = request.POST.get('zf')
    textlist.append(data65)
    textlist.append(data64)
    data66 = request.POST.get('qm1')
    textlist.append(data66)
    data67 = request.POST.get('qm2')
    textlist.append(data67)
    data68 = request.POST.get('qm3')
    textlist.append(data68)

    lx = request.POST.get('data06')
    if lx == '机场复习培训':
        if (data50 + data56 + data41 + data44 + data47 + data53 + data58) != '0000000':
            a = models.jichangfxkhb()
            a.frontdata1 = datan
            a.frontdata2 = datas
            a.frontdata3 = request.POST.get('data06')
            a.frontdata4 = data50
            a.frontdata5 = data56
            a.frontdata6 = data41
            a.frontdata7 = data44
            a.frontdata8 = data47
            a.frontdata9 = data53
            a.frontdata10 = data58

            a.is_active = 0
            a.save()

        b = info_02.objects.get(data_01=datan)
        b.data_02 = round((float(b.data_02) - float(data50) / 10), 2)
        b.data_03 = round((float(b.data_03) - float(data56) / 10), 2)
        b.data_04 = round((float(b.data_04) - float(data41) / 10), 2)
        b.data_05 = round((float(b.data_05) - float(data44) / 10), 2)
        b.data_06 = round((float(b.data_06) - float(data47) / 10), 2)
        b.data_07 = round((float(b.data_07) - float(data53) / 10), 2)
        b.data_08 = round((float(b.data_08) - float(data58) / 10), 2)
        b.is_active = 0
        b.save()

        doc = Document("E:\pythonproject\ATC_Data\考核表\复训考核表\机场复训考核表\机场复训考核表样表.docx")

        count = 0

        for n in range(0, 20):
            table = doc.tables[n]
            for row in table.rows:
                for cell in row.cells:
                    if 'XXXX' in cell.text:
                        cell.text = cell.text.replace("XXXX", textlist[count])
                        count += 1
            n += 1

        date = time.strftime('%y.%m.%d')
        doc.save('E:\pythonproject\ATC_Data\考核表\复训考核表\机场复训考核表\%s%s机场复训考核表.docx' % (date, datan))
    if lx == '机场资格培训':
        if (data50 + data56 + data41 + data44 + data47 + data53 + data58) != '0000000':
            a = models.jichangzgkhb()
            a.frontdata1 = datan
            a.frontdata2 = datas
            a.frontdata3 = request.POST.get('data06')
            a.frontdata4 = data50
            a.frontdata5 = data56
            a.frontdata6 = data41
            a.frontdata7 = data44
            a.frontdata8 = data47
            a.frontdata9 = data53
            a.frontdata10 = data58

            a.is_active = 0
            a.save()

        b = info_02.objects.get(data_01=datan)
        b.data_02 = round((float(b.data_02) - float(data50) / 10), 2)
        b.data_03 = round((float(b.data_03) - float(data56) / 10), 2)
        b.data_04 = round((float(b.data_04) - float(data41) / 10), 2)
        b.data_05 = round((float(b.data_05) - float(data44) / 10), 2)
        b.data_06 = round((float(b.data_06) - float(data47) / 10), 2)
        b.data_07 = round((float(b.data_07) - float(data53) / 10), 2)
        b.data_08 = round((float(b.data_08) - float(data58) / 10), 2)
        b.is_active = 0
        b.save()

        doc = Document("E:\pythonproject\ATC_Data\考核表\资格考核表\机场资格考核表\机场资格考核表样表.docx")

        count = 0

        for n in range(0, 20):
            table = doc.tables[n]
            for row in table.rows:
                for cell in row.cells:
                    if 'XXXX' in cell.text:
                        cell.text = cell.text.replace("XXXX", textlist[count])
                        count += 1
            n += 1

        date = time.strftime('%y.%m.%d')
        doc.save('E:\pythonproject\ATC_Data\考核表\资格考核表\机场资格考核表\%s%s机场资格考核表.docx' % (date, datan))

    return render(request, 'ok.html')


def addACCfxkhb(request):  # 获取增加区调复训考核表记录的数据并保存

    textlist = []
    listnum = [11, 11, 2, 1, 1, 1, 2, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 2, 2, 1, 1, 1, 1, 1, 1, 1]

    datan = request.POST.get('data01')
    datas = request.POST.get('data02')

    for i in range(1, 8):
        name = 'data0' + str(i)
        a = request.POST.get(name)
        textlist.append(a)

    data08 = request.POST.get('data08')
    if data08 == '0':
        data08 = '初级'
    if data08 == '1':
        data08 = '中级'
    if data08 == '2':
        data08 = '高级'
    textlist.append(data08)
    data09 = request.POST.get('data09')
    textlist.append(data09)

    for i in range(10, 40):
        name = 'data' + str(i)
        a = request.POST.get(name)
        b = int(a)
        c = b * listnum[i - 10]
        b = str(b)
        c = str(c)
        textlist.append(b)
        textlist.append(c)

    data40 = request.POST.get('jgcs')
    data41 = request.POST.get('jgkf')
    data42 = request.POST.get('datajgz')
    data43 = request.POST.get('xtcs')
    data44 = request.POST.get('xtkf')
    data45 = request.POST.get('dataxtz')
    data46 = request.POST.get('yscs')
    data47 = request.POST.get('yskf')
    data48 = request.POST.get('dataysz')
    data49 = request.POST.get('lkcs')
    data50 = request.POST.get('lkkf')
    data51 = request.POST.get('datalkz')
    data52 = request.POST.get('tqcs')
    data53 = request.POST.get('tqkf')
    data54 = request.POST.get('datatqz')
    data55 = request.POST.get('ldcs')
    data56 = request.POST.get('ldkf')
    data57 = request.POST.get('dataldz')
    data58 = request.POST.get('qtkf')
    data59 = request.POST.get('data41')
    if data59 == '':
        data59 = "无"

    textlist.append(data59)
    dataqtkf = request.POST.get('data40')
    textlist.append(dataqtkf)

    data61 = request.POST.get('qkjl')
    textlist.append(data61)
    data62 = request.POST.get('zwpj')
    textlist.append(data62)
    data63 = request.POST.get('jypy')
    textlist.append(data63)
    data64 = request.POST.get('jg')
    if data64 == '0':
        data64 = '通过'
    if data64 == '1':
        data64 = '不通过'
    if data64 == '2':
        data64 = '待定'
    data65 = request.POST.get('zf')
    textlist.append(data65)
    textlist.append(data64)
    data66 = request.POST.get('qm1')
    textlist.append(data66)
    data67 = request.POST.get('qm2')
    textlist.append(data67)
    data68 = request.POST.get('qm3')
    textlist.append(data68)

    lx = request.POST.get('data06')
    if lx == '区调复习培训':
        if (data50 + data56 + data41 + data44 + data47 + data53 + data58) != '0000000':
            a = models.ACCfxkhb()
            a.frontdata1 = datan
            a.frontdata2 = datas
            a.frontdata3 = request.POST.get('data06')
            a.frontdata4 = data50
            a.frontdata5 = data56
            a.frontdata6 = data41
            a.frontdata7 = data44
            a.frontdata8 = data47
            a.frontdata9 = data53
            a.frontdata10 = data58

            a.is_active = 0
            a.save()

        b = info_02.objects.get(data_01=datan)
        b.data_02 = round((float(b.data_02) - float(data50) / 10), 2)
        b.data_03 = round((float(b.data_03) - float(data56) / 10), 2)
        b.data_04 = round((float(b.data_04) - float(data41) / 10), 2)
        b.data_05 = round((float(b.data_05) - float(data44) / 10), 2)
        b.data_06 = round((float(b.data_06) - float(data47) / 10), 2)
        b.data_07 = round((float(b.data_07) - float(data53) / 10), 2)
        b.data_08 = round((float(b.data_08) - float(data58) / 10), 2)
        b.is_active = 0
        b.save()

        doc = Document("E:\pythonproject\ATC_Data\考核表\复训考核表\区调复训考核表\区调复训考核表样表.docx")

        count = 0

        for n in range(0, 20):
            table = doc.tables[n]
            for row in table.rows:
                for cell in row.cells:
                    if 'XXXX' in cell.text:
                        cell.text = cell.text.replace("XXXX", textlist[count])
                        count += 1
            n += 1

        date = time.strftime('%y.%m.%d')
        doc.save('E:\pythonproject\ATC_Data\考核表\复训考核表\区调复训考核表\%s%s区调复训考核表.docx' % (date, datan))
    if lx == '区调资格培训':
        if (data50 + data56 + data41 + data44 + data47 + data53 + data58) != '0000000':
            a = models.ACCzgkhb()
            a.frontdata1 = datan
            a.frontdata2 = datas
            a.frontdata3 = request.POST.get('data06')
            a.frontdata4 = data50
            a.frontdata5 = data56
            a.frontdata6 = data41
            a.frontdata7 = data44
            a.frontdata8 = data47
            a.frontdata9 = data53
            a.frontdata10 = data58

            a.is_active = 0
            a.save()

        b = info_02.objects.get(data_01=datan)
        b.data_02 = round((float(b.data_02) - float(data50) / 10), 2)
        b.data_03 = round((float(b.data_03) - float(data56) / 10), 2)
        b.data_04 = round((float(b.data_04) - float(data41) / 10), 2)
        b.data_05 = round((float(b.data_05) - float(data44) / 10), 2)
        b.data_06 = round((float(b.data_06) - float(data47) / 10), 2)
        b.data_07 = round((float(b.data_07) - float(data53) / 10), 2)
        b.data_08 = round((float(b.data_08) - float(data58) / 10), 2)
        b.is_active = 0
        b.save()

        doc = Document("E:\pythonproject\ATC_Data\考核表\资格考核表\区调资格考核表\区调资格考核表样表.docx")

        count = 0

        for n in range(0, 20):
            table = doc.tables[n]
            for row in table.rows:
                for cell in row.cells:
                    if 'XXXX' in cell.text:
                        cell.text = cell.text.replace("XXXX", textlist[count])
                        count += 1
            n += 1

        date = time.strftime('%y.%m.%d')
        doc.save('E:\pythonproject\ATC_Data\考核表\资格考核表\区调资格考核表\%s%s区调资格考核表.docx' % (date, datan))

    return render(request, 'ok.html')

def addAPPfxkhb(request):  # 获取增加进近复训考核表记录的数据并保存

    textlist = []
    listnum = [11, 1, 1, 1, 1, 2, 2, 1, 2, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 2, 1,1, 1,1,1,1 ]

    datan = request.POST.get('data01')
    datas = request.POST.get('data02')

    for i in range(1, 8):
        name = 'data0' + str(i)
        a = request.POST.get(name)
        textlist.append(a)

    data08 = request.POST.get('data08')
    if data08 == '0':
        data08 = '初级'
    if data08 == '1':
        data08 = '中级'
    if data08 == '2':
        data08 = '高级'
    textlist.append(data08)
    data09 = request.POST.get('data09')
    textlist.append(data09)

    for i in range(10, 39):
        name = 'data' + str(i)
        a = request.POST.get(name)
        b = int(a)
        c = b * listnum[i - 10]
        b = str(b)
        c = str(c)
        textlist.append(b)
        textlist.append(c)

    data40 = request.POST.get('jgcs')
    data41 = request.POST.get('jgkf')
    data42 = request.POST.get('datajgz')
    data43 = request.POST.get('xtcs')
    data44 = request.POST.get('xtkf')
    data45 = request.POST.get('dataxtz')
    data46 = request.POST.get('yscs')
    data47 = request.POST.get('yskf')
    data48 = request.POST.get('dataysz')
    data49 = request.POST.get('lkcs')
    data50 = request.POST.get('lkkf')
    data51 = request.POST.get('datalkz')
    data52 = request.POST.get('tqcs')
    data53 = request.POST.get('tqkf')
    data54 = request.POST.get('datatqz')
    data55 = request.POST.get('ldcs')
    data56 = request.POST.get('ldkf')
    data57 = request.POST.get('dataldz')
    data58 = request.POST.get('qtkf')
    data59 = request.POST.get('data41')
    if data59 == '':
        data59 = "无"

    textlist.append(data59)
    dataqtkf = request.POST.get('data40')
    textlist.append(dataqtkf)

    data61 = request.POST.get('qkjl')
    textlist.append(data61)
    data62 = request.POST.get('zwpj')
    textlist.append(data62)
    data63 = request.POST.get('jypy')
    textlist.append(data63)
    data64 = request.POST.get('jg')
    if data64 == '0':
        data64 = '通过'
    if data64 == '1':
        data64 = '不通过'
    if data64 == '2':
        data64 = '待定'
    data65 = request.POST.get('zf')
    textlist.append(data65)
    textlist.append(data64)
    data66 = request.POST.get('qm1')
    textlist.append(data66)
    data67 = request.POST.get('qm2')
    textlist.append(data67)
    data68 = request.POST.get('qm3')
    textlist.append(data68)

    lx = request.POST.get('data06')
    if lx == '进近复习培训':
        if (data50 + data56 + data41 + data44 + data47 + data53 + data58) != '0000000':
            a = models.APPfxkhb()
            a.frontdata1 = datan
            a.frontdata2 = datas
            a.frontdata3 = request.POST.get('data06')
            a.frontdata4 = data50
            a.frontdata5 = data56
            a.frontdata6 = data41
            a.frontdata7 = data44
            a.frontdata8 = data47
            a.frontdata9 = data53
            a.frontdata10 = data58

            a.is_active = 0
            a.save()

        b = info_02.objects.get(data_01=datan)
        b.data_02 = round((float(b.data_02) - float(data50) / 10), 2)
        b.data_03 = round((float(b.data_03) - float(data56) / 10), 2)
        b.data_04 = round((float(b.data_04) - float(data41) / 10), 2)
        b.data_05 = round((float(b.data_05) - float(data44) / 10), 2)
        b.data_06 = round((float(b.data_06) - float(data47) / 10), 2)
        b.data_07 = round((float(b.data_07) - float(data53) / 10), 2)
        b.data_08 = round((float(b.data_08) - float(data58) / 10), 2)
        b.is_active = 0
        b.save()

        doc = Document("E:\pythonproject\ATC_Data\考核表\复训考核表\进近复训考核表\进近复训考核表样表.docx")

        count = 0

        for n in range(0, 20):
            table = doc.tables[n]
            for row in table.rows:
                for cell in row.cells:
                    if 'XXXX' in cell.text:
                        cell.text = cell.text.replace("XXXX", textlist[count])
                        count += 1
            n += 1

        date = time.strftime('%y.%m.%d')
        doc.save('E:\pythonproject\ATC_Data\考核表\复训考核表\进近复训考核表\%s%s进近复训考核表.docx' % (date, datan))

    if lx == '进近资格培训':
        if (data50 + data56 + data41 + data44 + data47 + data53 + data58) != '0000000':
            a = models.APPzgkhb()
            a.frontdata1 = datan
            a.frontdata2 = datas
            a.frontdata3 = request.POST.get('data06')
            a.frontdata4 = data50
            a.frontdata5 = data56
            a.frontdata6 = data41
            a.frontdata7 = data44
            a.frontdata8 = data47
            a.frontdata9 = data53
            a.frontdata10 = data58

            a.is_active = 0
            a.save()

        b = info_02.objects.get(data_01=datan)
        b.data_02 = round((float(b.data_02) - float(data50) / 10), 2)
        b.data_03 = round((float(b.data_03) - float(data56) / 10), 2)
        b.data_04 = round((float(b.data_04) - float(data41) / 10), 2)
        b.data_05 = round((float(b.data_05) - float(data44) / 10), 2)
        b.data_06 = round((float(b.data_06) - float(data47) / 10), 2)
        b.data_07 = round((float(b.data_07) - float(data53) / 10), 2)
        b.data_08 = round((float(b.data_08) - float(data58) / 10), 2)
        b.is_active = 0
        b.save()

        doc = Document("E:\pythonproject\ATC_Data\考核表\资格考核表\进近资格考核表\进近资格考核表样表.docx")

        count = 0

        for n in range(0, 20):
            table = doc.tables[n]
            for row in table.rows:
                for cell in row.cells:
                    if 'XXXX' in cell.text:
                        cell.text = cell.text.replace("XXXX", textlist[count])
                        count += 1
            n += 1

        date = time.strftime('%y.%m.%d')
        doc.save('E:\pythonproject\ATC_Data\考核表\资格考核表\进近资格考核表\%s%s进近资格考核表.docx' % (date, datan))

    return render(request, 'ok.html')

def addACCgqkhb(request):  # 获取增加区调岗前考核表记录的数据并保存

    textlist = []
    listnum=[2,10,2,2,3,2,3,2,2,2,2,2,5,2,2,2,2,1,2,2,5,1,1,1,1,2,1,2,2,2,1,2,2,1,1,1,2,10,1]

    datan = request.POST.get('data01')
    datas = request.POST.get('data02')

    for i in range(1,6):
        name = 'data0' + str(i)
        a=request.POST.get(name)

        textlist.append(a)

    data06 = request.POST.get('data06')
    if data06 == '0':
        data06 = '初级'
    if data06 == '1':
        data06 = '中级'
    if data06 == '2':
        data06 = '高级'
    textlist.append(data06)

    data07 = request.POST.get('data07')
    textlist.append(data07)
    nr = request.POST.get('nr')
    textlist.append(nr)
    data08 = request.POST.get('data08')
    textlist.append(data08)
    c = str(int(data08) * 2)
    textlist.append(c)
    data09 = request.POST.get('data09')
    textlist.append(data09)
    c = str(int(data09) * 10)
    textlist.append(c)

    for i in range(10,47):
        name = 'data' + str(i)
        a=request.POST.get(name)
        b=int(a)
        c = b * listnum[i-8]
        b=str(b)
        c=str(c)
        textlist.append(b)
        textlist.append(c)

    data47 = request.POST.get('jgcs')
    data48 = request.POST.get('jgkf')
    data49 = request.POST.get('datajgz')
    data50 = request.POST.get('xtcs')
    data51 = request.POST.get('xtkf')
    data52 = request.POST.get('dataxtz')
    data53 = request.POST.get('yscs')
    data54 = request.POST.get('yskf')
    data55 = request.POST.get('dataysz')
    data56 = request.POST.get('lkcs')
    data57 = request.POST.get('lkkf')
    data58 = request.POST.get('datalkz')
    data59 = request.POST.get('tqcs')
    data60 = request.POST.get('tqkf')
    data61 = request.POST.get('datatqz')
    data62 = request.POST.get('ldcs')
    data63 = request.POST.get('ldkf')
    data64 = request.POST.get('dataldz')
    data65 = request.POST.get('qtkf')
    data66 = request.POST.get('data48')
    if data66=='':
       data66="无"

    textlist.append(data66)
    dataqtkf = request.POST.get('data47')
    textlist.append(dataqtkf)

    data67 = request.POST.get('qkjl')
    textlist.append(data67)
    data68 = request.POST.get('zwpj')
    textlist.append(data68)
    data69 = request.POST.get('jypy')
    textlist.append(data69)
    data70 = request.POST.get('jg')
    if data70 == '0':
        data70 = '通过'
    if data70 == '1':
        data70 = '不通过'
    if data70 == '2':
        data70 = '待定'
    data71 = request.POST.get('zf')
    textlist.append(data71)
    textlist.append(data70)
    data72 = request.POST.get('qm1')
    textlist.append(data72)
    data73 = request.POST.get('qm2')
    textlist.append(data73)

    if (data57 + data63 + data48 + data51 + data54 + data60 + data65) != '0000000':
        a = models.ACCgqkhb()
        a.frontdata1 = datan
        a.frontdata2 = datas
        a.frontdata3 = request.POST.get('data07')
        a.frontdata4 = data57
        a.frontdata5 = data63
        a.frontdata6 = data48
        a.frontdata7 = data51
        a.frontdata8 = data54
        a.frontdata9 = data60
        a.frontdata10 = data65

        a.is_active = 0
        a.save()

    b = info_02.objects.get(data_01=datan)
    b.data_02 = round((float(b.data_02) - float(data57) / 10),2)
    b.data_03 = round((float(b.data_03) - float(data63) / 10),2)
    b.data_04 = round((float(b.data_04) - float(data48) / 10),2)
    b.data_05 = round((float(b.data_05) - float(data51) / 10),2)
    b.data_06 = round((float(b.data_06) - float(data54) / 10),2)
    b.data_07 = round((float(b.data_07) - float(data60) / 10),2)
    b.data_08 = round((float(b.data_08) - float(data65) / 10),2)
    b.is_active = 0
    b.save()

    doc = Document("E:\pythonproject\ATC_Data\考核表\岗前考核表\区调岗前考核表\区调岗前考核表样表.docx")

    count = 0

    for n in range(0, 20):
        table = doc.tables[n]
        for row in table.rows:
            for cell in row.cells:
                if 'XXXX' in cell.text:
                    cell.text = cell.text.replace("XXXX", textlist[count])
                    count += 1
        n += 1

    date = time.strftime('%y.%m.%d')
    doc.save('E:\pythonproject\ATC_Data\考核表\岗前考核表\区调岗前考核表\%s%s区调岗前考核表.docx' %(date,datan))
    return render(request, 'ok.html')

def addAPPgqkhb(request):  # 获取增加进近岗前考核表记录的数据并保存

    textlist = []
    listnum=[2,10,10,4,2,2,2,2,2,2,2,2,2,2,5,2,1,2,2,3,2,2,2,5,1,1,1,1,1,1,2,2,2,1,2,2,1,1,1,2,10,1]

    datan = request.POST.get('data01')
    datas = request.POST.get('data02')

    for i in range(1,6):
        name = 'data0' + str(i)
        a=request.POST.get(name)

        textlist.append(a)

    data06 = request.POST.get('data06')
    if data06 == '0':
        data06 = '初级'
    if data06 == '1':
        data06 = '中级'
    if data06 == '2':
        data06 = '高级'
    textlist.append(data06)

    data07 = request.POST.get('data07')
    textlist.append(data07)
    nr = request.POST.get('nr')
    textlist.append(nr)
    data08 = request.POST.get('data08')
    textlist.append(data08)
    c = str(int(data08) * 2)
    textlist.append(c)
    data09 = request.POST.get('data09')
    textlist.append(data09)
    c = str(int(data09) * 10)
    textlist.append(c)

    for i in range(10,50):
        name = 'data' + str(i)
        a=request.POST.get(name)
        b=int(a)
        c = b * listnum[i-8]
        b=str(b)
        c=str(c)
        textlist.append(b)
        textlist.append(c)

    data47 = request.POST.get('jgcs')
    data48 = request.POST.get('jgkf')
    data49 = request.POST.get('datajgz')
    data50 = request.POST.get('xtcs')
    data51 = request.POST.get('xtkf')
    data52 = request.POST.get('dataxtz')
    data53 = request.POST.get('yscs')
    data54 = request.POST.get('yskf')
    data55 = request.POST.get('dataysz')
    data56 = request.POST.get('lkcs')
    data57 = request.POST.get('lkkf')
    data58 = request.POST.get('datalkz')
    data59 = request.POST.get('tqcs')
    data60 = request.POST.get('tqkf')
    data61 = request.POST.get('datatqz')
    data62 = request.POST.get('ldcs')
    data63 = request.POST.get('ldkf')
    data64 = request.POST.get('dataldz')
    data65 = request.POST.get('qtkf')
    data66 = request.POST.get('data51')
    if data66=='':
       data66="无"

    textlist.append(data66)
    dataqtkf = request.POST.get('data50')
    textlist.append(dataqtkf)

    data67 = request.POST.get('qkjl')
    textlist.append(data67)
    data68 = request.POST.get('zwpj')
    textlist.append(data68)
    data69 = request.POST.get('jypy')
    textlist.append(data69)
    data70 = request.POST.get('jg')
    if data70 == '0':
        data70 = '通过'
    if data70 == '1':
        data70 = '不通过'
    if data70 == '2':
        data70 = '待定'
    data71 = request.POST.get('zf')
    textlist.append(data71)
    textlist.append(data70)
    data72 = request.POST.get('qm1')
    textlist.append(data72)
    data73 = request.POST.get('qm2')
    textlist.append(data73)

    if (data57 + data63 + data48 + data51 + data54 + data60 + data65) != '0000000':
        a = models.APPgqkhb()
        a.frontdata1 = datan
        a.frontdata2 = datas
        a.frontdata3 = request.POST.get('data07')
        a.frontdata4 = data57
        a.frontdata5 = data63
        a.frontdata6 = data48
        a.frontdata7 = data51
        a.frontdata8 = data54
        a.frontdata9 = data60
        a.frontdata10 = data65

        a.is_active = 0
        a.save()

    b = info_02.objects.get(data_01=datan)
    b.data_02 = round((float(b.data_02) - float(data57) / 10),2)
    b.data_03 = round((float(b.data_03) - float(data63) / 10),2)
    b.data_04 = round((float(b.data_04) - float(data48) / 10),2)
    b.data_05 = round((float(b.data_05) - float(data51) / 10),2)
    b.data_06 = round((float(b.data_06) - float(data54) / 10),2)
    b.data_07 = round((float(b.data_07) - float(data60) / 10),2)
    b.data_08 = round((float(b.data_08) - float(data65) / 10),2)
    b.is_active = 0
    b.save()

    doc = Document("E:\pythonproject\ATC_Data\考核表\岗前考核表\进近岗前考核表\进近岗前考核表样表.docx")

    count = 0

    for n in range(0, 20):
        table = doc.tables[n]
        for row in table.rows:
            for cell in row.cells:
                if 'XXXX' in cell.text:
                    cell.text = cell.text.replace("XXXX", textlist[count])
                    count += 1
        n += 1

    date = time.strftime('%y.%m.%d')
    doc.save('E:\pythonproject\ATC_Data\考核表\岗前考核表\进近岗前考核表\%s%s进近岗前考核表.docx' %(date,datan))
    return render(request, 'ok.html')


def addjichanggqkhb(request):  # 获取增加机场岗前考核表记录的数据并保存

    textlist = []
    listnum=[3,10,2,1,2,2,3,3,3,2,5,2,1,2,2,1,3,1,5,5,1,1,1,1,1,1,2,2,1,2,2,1,1,1,2,2]

    datan = request.POST.get('data01')
    datas = request.POST.get('data02')

    for i in range(1,6):
        name = 'data0' + str(i)
        a=request.POST.get(name)

        textlist.append(a)

    data06 = request.POST.get('data06')
    if data06 == '0':
        data06 = '初级'
    if data06 == '1':
        data06 = '中级'
    if data06 == '2':
        data06 = '高级'
    textlist.append(data06)

    data07 = request.POST.get('data07')
    textlist.append(data07)
    nr = request.POST.get('nr')
    textlist.append(nr)
    data08 = request.POST.get('data08')
    textlist.append(data08)
    c = str(int(data08) * 2)
    textlist.append(c)
    data09 = request.POST.get('data09')
    textlist.append(data09)
    c = str(int(data09) * 10)
    textlist.append(c)

    for i in range(10,44):
        name = 'data' + str(i)
        a=request.POST.get(name)
        b=int(a)
        c = b * listnum[i-8]
        b=str(b)
        c=str(c)
        textlist.append(b)
        textlist.append(c)

    data47 = request.POST.get('jgcs')
    data48 = request.POST.get('jgkf')
    data49 = request.POST.get('datajgz')
    data50 = request.POST.get('xtcs')
    data51 = request.POST.get('xtkf')
    data52 = request.POST.get('dataxtz')
    data53 = request.POST.get('yscs')
    data54 = request.POST.get('yskf')
    data55 = request.POST.get('dataysz')
    data56 = request.POST.get('lkcs')
    data57 = request.POST.get('lkkf')
    data58 = request.POST.get('datalkz')
    data59 = request.POST.get('tqcs')
    data60 = request.POST.get('tqkf')
    data61 = request.POST.get('datatqz')
    data62 = request.POST.get('ldcs')
    data63 = request.POST.get('ldkf')
    data64 = request.POST.get('dataldz')
    data65 = request.POST.get('qtkf')
    data66 = request.POST.get('data45')
    if data66=='':
       data66="无"

    textlist.append(data66)
    dataqtkf = request.POST.get('data44')
    textlist.append(dataqtkf)

    data67 = request.POST.get('qkjl')
    textlist.append(data67)
    data68 = request.POST.get('zwpj')
    textlist.append(data68)
    data69 = request.POST.get('jypy')
    textlist.append(data69)
    data70 = request.POST.get('jg')
    if data70 == '0':
        data70 = '通过'
    if data70 == '1':
        data70 = '不通过'
    if data70 == '2':
        data70 = '待定'
    data71 = request.POST.get('zf')
    textlist.append(data71)
    textlist.append(data70)
    data72 = request.POST.get('qm1')
    textlist.append(data72)
    data73 = request.POST.get('qm2')
    textlist.append(data73)

    if (data57 + data63 + data48 + data51 + data54 + data60 + data65) != '0000000':
        a = models.jichanggqkhb()
        a.frontdata1 = datan
        a.frontdata2 = datas
        a.frontdata3 = request.POST.get('data07')
        a.frontdata4 = data57
        a.frontdata5 = data63
        a.frontdata6 = data48
        a.frontdata7 = data51
        a.frontdata8 = data54
        a.frontdata9 = data60
        a.frontdata10 = data65

        a.is_active = 0
        a.save()

    b = info_02.objects.get(data_01=datan)
    b.data_02 = round((float(b.data_02) - float(data57) / 10),2)
    b.data_03 = round((float(b.data_03) - float(data63) / 10),2)
    b.data_04 = round((float(b.data_04) - float(data48) / 10),2)
    b.data_05 = round((float(b.data_05) - float(data51) / 10),2)
    b.data_06 = round((float(b.data_06) - float(data54) / 10),2)
    b.data_07 = round((float(b.data_07) - float(data60) / 10),2)
    b.data_08 = round((float(b.data_08) - float(data65) / 10),2)
    b.is_active = 0
    b.save()

    doc = Document("E:\pythonproject\ATC_Data\考核表\岗前考核表\机场岗前考核表\机场岗前考核表样表.docx")

    count = 0

    for n in range(0, 20):
        table = doc.tables[n]
        for row in table.rows:
            for cell in row.cells:
                if 'XXXX' in cell.text:
                    cell.text = cell.text.replace("XXXX", textlist[count])
                    count += 1
        n += 1

    date = time.strftime('%y.%m.%d')
    doc.save('E:\pythonproject\ATC_Data\考核表\岗前考核表\机场岗前考核表\%s%s机场岗前考核表.docx' %(date,datan))
    return render(request, 'ok.html')

def getkfdetail(request):
    b = jianchabaogao.objects.get(data_01=姓名)#从检查报告读出记录
    #筛出frontdata7不为表扬和正常的项（技能缺陷，现场纠违，无后果违章）
    #从这个人所在部门（app/acc）对应的表里（zgkhb/fxkhb/gqkhb）搜出名字对应的记录
    #如果是塔台的还要从jichang3张表读

